from docling.document_converter import DocumentConverter
import os
import logging
import sys
import tempfile
import shutil

def extract_document(file_path, output_path="extract.txt", disable_symlinks_warning=True, use_simple_extraction=False):
    """
    Extract text content from a document using docling and save it to a text file.
    
    Parameters:
    -----------
    file_path : str
        Path to the document file to extract (PDF, DOCX, TXT, etc.)
    output_path : str, optional
        Path where the extracted text will be saved, defaults to 'extract.txt'
    disable_symlinks_warning : bool, optional
        Whether to disable HuggingFace symlinks warning, defaults to True
    use_simple_extraction : bool, optional
        If True, uses a simpler approach that may work better on Windows without admin rights
        
    Returns:
    --------
    bool
        True if extraction was successful, False otherwise
    
    Example:
    --------
    >>> extract_document("path/to/document.pdf")
    >>> extract_document("path/to/document.docx", "custom_output.txt")
    """
    # Disable HuggingFace symlinks warning if requested
    if disable_symlinks_warning:
        os.environ["HF_HUB_DISABLE_SYMLINKS_WARNING"] = "1"
    
    try:
        # Check if the file exists
        if not os.path.exists(file_path):
            logging.error(f"File not found: {file_path}")
            return False
        
        if use_simple_extraction:
            # Simpler approach for Windows environments without symlink privileges
            text = _extract_document_simple(file_path)
            
            # Save the extracted text to the output file
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
                
            logging.info(f"Document extracted using simple method and saved to {output_path}")
            return True
        else:
            # Standard approach using docling's full capabilities
            # Initialize the document converter
            converter = DocumentConverter()
            
            # Extract text from the document
            result = converter.convert(file_path)
            
            # Export the document to markdown text
            text = result.document.export_to_markdown()
            
            # Save the extracted text to the output file
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
            
            logging.info(f"Document extracted successfully and saved to {output_path}")
            return True
    
    except Exception as e:
        logging.error(f"Error extracting document: {str(e)}")
        
        # If standard extraction fails, try the simple approach as fallback
        if not use_simple_extraction:
            logging.info("Attempting fallback to simple extraction method...")
            return extract_document(file_path, output_path, disable_symlinks_warning, True)
        
        return False

def _extract_document_simple(file_path):
    """
    A simpler approach to extract text from documents that doesn't rely on complex models.
    This may work better in environments where the full docling functionality has permission issues.
    
    Parameters:
    -----------
    file_path : str
        Path to the document file to extract
        
    Returns:
    --------
    str
        Extracted text content
    """
    import mimetypes
    file_type = mimetypes.guess_type(file_path)[0]
    
    text = ""
    
    # Handle PDF files
    if file_type == 'application/pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text += page.get_text()
            doc.close()
        except ImportError:
            logging.warning("PyMuPDF not available, falling back to docling CLI")
            text = _try_docling_cli(file_path)
    
    # Handle DOCX files
    elif file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        try:
            from docx import Document
            doc = Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():  # Skip empty paragraphs
                    paragraphs.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():  # Skip empty cells
                            row_text.append(cell.text.strip())
                    if row_text:  # Skip empty rows
                        paragraphs.append(" | ".join(row_text))
            
            # Join all extracted text
            text = "\n\n".join(paragraphs)
        except ImportError:
            logging.warning("python-docx not available, falling back to docling CLI")
            text = _try_docling_cli(file_path)
        except Exception as e:
            logging.warning(f"Error extracting text from DOCX using python-docx: {str(e)}")
            logging.warning("Falling back to docling CLI")
            text = _try_docling_cli(file_path)
    
    # Handle TXT files
    elif file_type == 'text/plain':
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
    
    # For other file types or if type detection failed, try docling CLI
    else:
        text = _try_docling_cli(file_path)
    
    return text

def _try_docling_cli(file_path):
    """Try to use the docling CLI to extract text."""
    import subprocess
    import tempfile
    
    temp_output = tempfile.NamedTemporaryFile(delete=False, suffix='.md')
    temp_output.close()
    
    try:
        result = subprocess.run(
            ["docling", file_path, "--output", temp_output.name],
            capture_output=True, 
            text=True, 
            check=False
        )
        
        if result.returncode == 0:
            with open(temp_output.name, 'r', encoding='utf-8') as f:
                return f.read()
        else:
            logging.error(f"docling CLI failed: {result.stderr}")
            return f"[Failed to extract content]\nError: {result.stderr}"
    except Exception as e:
        logging.error(f"Error using docling CLI: {str(e)}")
        return f"[Failed to extract content]\nError: {str(e)}"
    finally:
        # Clean up temp file
        try:
            os.unlink(temp_output.name)
        except:
            pass
    

if __name__ == "__main__":
    # Configure logging
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s'
    )
    
    import sys
    
    if len(sys.argv) > 1:
        # If file path is provided as command line argument
        file_path = sys.argv[1]
        output_path = "extracted_text.txt"
        if len(sys.argv) > 2:
            output_path = sys.argv[2]
        
        logging.info(f"Extracting document: {file_path}")
        logging.info(f"Output path: {output_path}")
        
        # Try with the simple extraction method first to avoid Windows permission issues
        result = extract_document(
            file_path, 
            output_path,
            disable_symlinks_warning=True,
            use_simple_extraction=True
        )
        
        if result:
            logging.info(f"Document extracted successfully: {output_path}")
        else:
            logging.error(f"Failed to extract document: {file_path}")
    else:
        # Default file path
        # Try with the simple extraction method first to avoid Windows permission issues
        extract_document(
            r"D:\Coding\Context-cache-system\Resume.docx", 
            "extracted_text.txt",
            disable_symlinks_warning=True,
            use_simple_extraction=True
        )
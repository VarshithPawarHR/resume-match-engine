import sys
import os

# Add the project root to the Python path to enable running this script directly
project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
if project_root not in sys.path:
    sys.path.insert(0, project_root)

from google import genai
from google.genai import types
from pathlib import Path
import io
import os
import sys
import time
import json
from datetime import datetime
import httpx
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type

sys.path.append(os.path.dirname(os.path.abspath(__file__)))
from utils.config import GEMINI_API_KEY, GEMINI_MODEL
from utils import prompts
import concurrent.futures
import threading
# Import the SQLite database manager
from utils import db_manager

# Define the structured output schema for ATS scoring
ATS_SCHEMA = {
    "type": "object",
    "properties": {
        "candidate_name": {"type": "string"},
        "position_applied": {"type": "string"},
        "company": {"type": "string"},
        "overall_fit_score": {"type": "number", "minimum": 0, "maximum": 100},
        "recommendation": {"type": "string", "enum": ["APPROVED", "REJECTED"]},
        "fit_level": {"type": "string", "enum": ["HIGH_FIT", "MEDIUM_FIT", "LOW_FIT", "NO_FIT"]},
        "key_strengths": {
            "type": "array",
            "items": {"type": "string"}
        },
        "major_concerns": {
            "type": "array",
            "items": {"type": "string"}
        },
        "skills_assessment": {
            "type": "object",
            "properties": {
                "required_skills_match": {"type": "integer", "minimum": 0, "maximum": 100},
                "preferred_skills_match": {"type": "integer", "minimum": 0, "maximum": 100},
                "critical_skills_missing": {
                    "type": "array",
                    "items": {"type": "string"}
                },
                "skill_gaps_impact": {"type": "string", "enum": ["Low", "Medium", "High", "Critical"]}
            },
            "required": ["required_skills_match", "preferred_skills_match", "critical_skills_missing", "skill_gaps_impact"]
        },
        "experience_fit": {
            "type": "object",
            "properties": {
                "years_required": {"type": "number", "minimum": 0},
                "years_candidate_has": {"type": "number", "minimum": 0},
                "experience_relevance": {"type": "string", "enum": ["High", "Medium", "Low", "None"]},
                "project_quality": {"type": "string", "enum": ["Excellent", "Good", "Average", "Poor"]}
            },
            "required": ["years_required", "years_candidate_has", "experience_relevance", "project_quality"]
        },
        "hiring_decision_factors": {
            "type": "object",
            "properties": {
                "technical_competency": {"type": "integer", "minimum": 0, "maximum": 100},
                "experience_level": {"type": "integer", "minimum": 0, "maximum": 100},
                "cultural_fit_indicators": {"type": "integer", "minimum": 0, "maximum": 100},
                "growth_potential": {"type": "integer", "minimum": 0, "maximum": 100},
                "immediate_productivity": {"type": "integer", "minimum": 0, "maximum": 100}
            },
            "required": ["technical_competency", "experience_level", "cultural_fit_indicators", "growth_potential", "immediate_productivity"]
        }
    },
    "required": [
        "candidate_name", "position_applied", "company", "overall_fit_score", 
        "recommendation", "fit_level", "key_strengths", 
        "major_concerns", "skills_assessment", "experience_fit", "hiring_decision_factors"
    ]
}


@retry(
    stop=stop_after_attempt(3),
    wait=wait_exponential(multiplier=1, min=4, max=10),
    retry=retry_if_exception_type((httpx.ConnectError, ConnectionError, TimeoutError))
)
def upload_file_with_retry(client, file_io, mime_type, username=None, filename=None, file_type=None):
    """Upload file with retry logic for connection issues and store info in database"""
    try:
        uploaded = client.files.upload(file=file_io, config=dict(mime_type=mime_type))
        file_id = None
        # Store file information in the database if username is provided
        if username and filename and file_type:
            # Get or create user (now returns user_id as string)
            user = db_manager.get_or_create_user(username)

            # Save file record
            file_path = str(Path(filename).absolute()) if isinstance(filename, (str, Path)) else "memory_file"
            file_id = db_manager.save_file_record(
                user_id=user['id'],  # user['id'] is now the username string
                filename=filename,
                file_path=file_path,
                file_type=file_type,
                mime_type=mime_type,
                gemini_file_id=uploaded.name
            )

        return uploaded, file_id
    except Exception as e:
        print(f"Upload attempt failed: {e}")
        raise

def analyze_two_files(file1_path, file2_path, username=None, user_query="", prompt_type="analysis", use_structured_output=True):
    """
    Analyze two files (PDF or DOCX) with Gemini in one simple function.

    Args:
        file1_path (str): Path to the first file (.pdf or .docx)
        file2_path (str): Path to the second file (.pdf or .docx)
        username (str): Username to associate with the files and analysis results
        user_query (str): The user's analysis/query prompt
        prompt_type (str): Type of system prompt to use (default: "analysis")
        use_structured_output (bool): Whether to use structured output with ATS schema

    Returns:
        str: The AI-generated analysis result (JSON string if structured output is used)
    """
    # Helpers inside the function
    def _validate_file_type(path: Path):
        if path.suffix.lower() not in {".pdf", ".docx"}:
            raise ValueError(f"Unsupported file type: {path.suffix}. Only PDF and DOCX are supported.")

    def _get_mime_type(path: Path) -> str:
        return "application/pdf" if path.suffix.lower() == ".pdf" else "text/plain"

    def _convert_docx_to_text(path: Path) -> str:
        from docx import Document
        doc = Document(path)
        return "\n".join(p.text for p in doc.paragraphs)

    # Initialize Gemini client
    client = genai.Client(api_key=GEMINI_API_KEY)

    # Get or create user if username is provided
    user_id = None
    if username:
        user = db_manager.get_or_create_user(username)
        user_id = user['id']

    # Prepare both files
    uploaded_files = []
    db_file_ids = []
    
    for idx, f in enumerate([file1_path, file2_path]):
        path = Path(f)
        _validate_file_type(path)
        file_type = path.suffix.lower()[1:]  # Remove the dot
        
        if path.suffix.lower() == ".docx":
            text_content = _convert_docx_to_text(path)
            file_io = io.BytesIO(text_content.encode("utf-8"))
            mime_type = "text/plain"
        else:
            file_io = io.BytesIO(path.read_bytes())
            mime_type = _get_mime_type(path)
        
        try:
            uploaded, db_file_id = upload_file_with_retry(
                client, 
                file_io, 
                mime_type,
                username=username,
                filename=path.name,
                file_type=file_type
            )
        except Exception as e:
            print(f"Failed to upload {f} after retries: {e}")
            # Fallback: try to continue with text content if possible
            if path.suffix.lower() == ".pdf":
                print("PDF upload failed. Please check your network connection.")
                raise
            else:
                print("Proceeding with text content...")
                continue
                
        # Wait if processing
        while hasattr(uploaded, "state") and uploaded.state.name == "PROCESSING":
            time.sleep(2)
            uploaded = client.files.get(name=uploaded.name)
            
        uploaded_files.append(uploaded)
        
        # Store the database file ID if available
        if db_file_id:
            db_file_ids.append(db_file_id)
        else:
            db_file_ids.append(None)

    # Create cache
    display_name = f"Cache with {Path(file1_path).name} and {Path(file2_path).name}"
    cache = client.caches.create(
        model=GEMINI_MODEL,
        config=types.CreateCachedContentConfig(
            display_name=display_name,
            system_instruction={
                "analysis": prompts.system_prompt,
            }.get(prompt_type, prompts.system_prompt),
            contents=uploaded_files,
            ttl="1800s",
        ),
    )
    
    # Store cache information in the database if username is provided
    cache_id = None
    if username and len(db_file_ids) == 2 and all(db_file_ids):
        jd_file_id, resume_file_id = db_file_ids
        cache_id = db_manager.save_cache_record(
            user_id=user_id,
            cache_name=cache.name,
            display_name=display_name,
            jd_file_id=jd_file_id,
            resume_file_id=resume_file_id,
            ttl=1800
        )

    # Generate content with structured output if requested
    if use_structured_output:
        # Use a simple instruction since the system prompt handles the detailed requirements
        query_content = user_query if user_query.strip() else "Analyze the resume against the job description."
        
        response = client.models.generate_content(
            model=GEMINI_MODEL,
            contents=query_content,
            config=types.GenerateContentConfig(
                cached_content=cache.name,
                response_schema=ATS_SCHEMA,
                response_mime_type="application/json"
            ),
        )
        
        # Parse the JSON response and add timestamps
        try:
            result_data = json.loads(response.text)
            
            # Add automatic timestamp for when evaluation was performed
            current_time = datetime.now().isoformat() + "Z"
            result_data["evaluation_timestamp"] = current_time
            
            result_json = json.dumps(result_data, indent=2)
            
            # Store the analysis result in the database if username is provided
            if username and cache_id and len(db_file_ids) == 2 and all(db_file_ids):
                jd_file_id, resume_file_id = db_file_ids
                db_manager.save_analysis_result(
                    user_id=user_id,
                    cache_id=cache_id,
                    jd_file_id=jd_file_id,
                    resume_file_id=resume_file_id,
                    result_json=result_json
                )
            
            return result_json
            
        except json.JSONDecodeError:
            # If JSON parsing fails, return original response
            return response.text
            
    else:
        response = client.models.generate_content(
            model=GEMINI_MODEL,
            contents=user_query,
            config=types.GenerateContentConfig(cached_content=cache.name),
        )
        result_text = response.text
        
        # Store the analysis result in the database if username is provided
        if username and cache_id and len(db_file_ids) == 2 and all(db_file_ids):
            jd_file_id, resume_file_id = db_file_ids
            db_manager.save_analysis_result(
                user_id=user_id,
                cache_id=cache_id,
                jd_file_id=jd_file_id,
                resume_file_id=resume_file_id,
                result_json=json.dumps({"text_result": result_text})
            )
        
        return result_text


def analyze_bulk_resumes(jd_file_path, resume_file_paths, username=None, use_structured_output=True, batch_method="inline"):
    """
    Analyze multiple resumes against a single job description using Batch API.
    
    Args:
        jd_file_path (str): Path to the job description file (.pdf or .docx)
        resume_file_paths (list): List of paths to resume files (.pdf or .docx)
        username (str): Username to associate with the files and analysis results
        use_structured_output (bool): Whether to use structured output with ATS schema
        batch_method (str): "inline" for small batches (<20MB), "file" for large batches
        
    Returns:
        dict: Dictionary with resume paths as keys and analysis results as values
    """
    # Helper functions (same as in analyze_two_files)
    def _validate_file_type(path: Path):
        if path.suffix.lower() not in {".pdf", ".docx"}:
            raise ValueError(f"Unsupported file type: {path.suffix}. Only PDF and DOCX are supported.")

    def _get_mime_type(path: Path) -> str:
        return "application/pdf" if path.suffix.lower() == ".pdf" else "text/plain"

    def _convert_docx_to_text(path: Path) -> str:
        from docx import Document
        doc = Document(path)
        return "\n".join(p.text for p in doc.paragraphs)

    # Initialize Gemini client
    client = genai.Client(api_key=GEMINI_API_KEY)
    
    # Get or create user if username is provided
    user_id = None
    if username:
        user = db_manager.get_or_create_user(username)
        user_id = user['id']
    
    # Upload JD file first (this will be reused for all requests)
    print("Uploading job description...")
    jd_path = Path(jd_file_path)
    _validate_file_type(jd_path)
    file_type = jd_path.suffix.lower()[1:]  # Remove the dot
    
    if jd_path.suffix.lower() == ".docx":
        jd_text_content = _convert_docx_to_text(jd_path)
        jd_file_io = io.BytesIO(jd_text_content.encode("utf-8"))
        jd_mime_type = "text/plain"
    else:
        jd_file_io = io.BytesIO(jd_path.read_bytes())
        jd_mime_type = _get_mime_type(jd_path)
    
    jd_uploaded, jd_db_file_id = upload_file_with_retry(
        client, 
        jd_file_io, 
        jd_mime_type,
        username=username,
        filename=jd_path.name,
        file_type=file_type
    )
    
    # Wait if processing
    while hasattr(jd_uploaded, "state") and jd_uploaded.state.name == "PROCESSING":
        time.sleep(2)
        jd_uploaded = client.files.get(name=jd_uploaded.name)
    
    print(f"Job description uploaded: {jd_uploaded.name}")
    
    # Upload all resume files
    print(f"Uploading {len(resume_file_paths)} resumes...")
    uploaded_resumes = {}
    
    for resume_path in resume_file_paths:
        path = Path(resume_path)
        _validate_file_type(path)
        file_type = path.suffix.lower()[1:]  # Remove the dot
        
        if path.suffix.lower() == ".docx":
            text_content = _convert_docx_to_text(path)
            file_io = io.BytesIO(text_content.encode("utf-8"))
            mime_type = "text/plain"
        else:
            file_io = io.BytesIO(path.read_bytes())
            mime_type = _get_mime_type(path)
        
        try:
            uploaded, resume_db_file_id = upload_file_with_retry(
                client, 
                file_io, 
                mime_type,
                username=username,
                filename=path.name,
                file_type=file_type
            )
            
            # Wait if processing
            while hasattr(uploaded, "state") and uploaded.state.name == "PROCESSING":
                time.sleep(2)
                uploaded = client.files.get(name=uploaded.name)
            
            uploaded_resumes[resume_path] = (uploaded, resume_db_file_id)
            print(f"Uploaded: {path.name}")
            
        except Exception as e:
            print(f"Failed to upload {resume_path}: {e}")
            continue
    
    # Create batch requests
    print("Creating batch requests...")
    batch_requests = []
    
    for i, (resume_path, (resume_uploaded, resume_db_file_id)) in enumerate(uploaded_resumes.items()):
        # Create the request content with system instruction included in the user message
        analysis_prompt = f"""
{prompts.system_prompt}

Please analyze the following resume against the job description and provide a comprehensive ATS evaluation.

Job Description: (see attached file)
Resume: (see attached file)
"""
        
        request_config = {
            'contents': [{
                'parts': [
                    {'text': analysis_prompt},
                    {'file_data': {'file_uri': jd_uploaded.name, 'mime_type': jd_uploaded.mime_type}},
                    {'file_data': {'file_uri': resume_uploaded.name, 'mime_type': resume_uploaded.mime_type}}
                ],
                'role': 'user'
            }]
        }
        
        if use_structured_output:
            request_config['config'] = {
                'response_mime_type': 'application/json',
                'response_schema': ATS_SCHEMA
            }
        
        if batch_method == "inline":
            batch_requests.append(request_config)
        else:  # file method
            batch_requests.append({
                'key': f'resume-{i}',
                'request': request_config
            })
    
    # Create batch job
    print(f"Creating batch job with {len(batch_requests)} requests...")
    batch_job_id = None
    
    try:
        if batch_method == "inline":
            # Inline method for smaller batches
            batch_job = client.batches.create(
                model=GEMINI_MODEL,
                src=batch_requests,
                config={
                    'display_name': f"Bulk Resume Analysis - {len(batch_requests)} resumes",
                }
            )
        else:
            # File method for larger batches
            import tempfile
            with tempfile.NamedTemporaryFile(mode='w', suffix='.jsonl', delete=False) as f:
                for req in batch_requests:
                    f.write(json.dumps(req) + '\n')
                batch_file_path = f.name
            
            # Upload batch file
            uploaded_batch_file = client.files.upload(
                file=batch_file_path,
                config=types.UploadFileConfig(
                    display_name='bulk-resume-batch-requests',
                    mime_type='application/jsonl'
                )
            )
            
            batch_job = client.batches.create(
                model=GEMINI_MODEL,
                src=uploaded_batch_file.name,
                config={
                    'display_name': f"Bulk Resume Analysis - {len(batch_requests)} resumes",
                }
            )
            
            # Clean up temp file
            import os
            os.unlink(batch_file_path)
        
        print(f"Batch job created: {batch_job.name}")
        
        # Store batch job information in the database if username is provided
        if username and user_id:
            batch_job_id = db_manager.save_batch_job(
                user_id=user_id,
                job_name=batch_job.name,
                job_state=batch_job.state.name,
                num_requests=len(batch_requests)
            )
    except Exception as e:
        print(f"Error creating batch job: {e}")
        # Print more detailed error information if available
        if hasattr(e, 'details'):
            print(f"Error details: {e.details}")
        return {resume_path: f"Batch processing error: {str(e)}" for resume_path in uploaded_resumes.keys()}
    
    # Monitor job status
    print("Monitoring batch job status...")
    completed_states = {
        'JOB_STATE_SUCCEEDED',
        'JOB_STATE_FAILED', 
        'JOB_STATE_CANCELLED',
        'JOB_STATE_EXPIRED'
    }
    
    try:
        while batch_job.state.name not in completed_states:
            print(f"Current state: {batch_job.state.name}")
            
            # Update job status in database if available
            if batch_job_id and username:
                db_manager.update_batch_job_status(username, batch_job_id, batch_job.state.name)
                
            time.sleep(30)  # Wait 30 seconds before polling again
            batch_job = client.batches.get(name=batch_job.name)
        
        print(f"Batch job finished with state: {batch_job.state.name}")
        
        # Update final job status in database
        if batch_job_id and username:
            db_manager.update_batch_job_status(
                username,
                batch_job_id, 
                batch_job.state.name,
                datetime.now().isoformat()
            )
    except Exception as e:
        print(f"Error monitoring batch job: {e}")
        
        # Update job status to error in database
        if batch_job_id and username:
            db_manager.update_batch_job_status(
                username,
                batch_job_id, 
                'JOB_STATE_ERROR',
                datetime.now().isoformat()
            )
            
        return {resume_path: f"Batch job monitoring error: {str(e)}" for resume_path in uploaded_resumes.keys()}
    
    # Retrieve results
    results = {}
    
    if batch_job.state.name == 'JOB_STATE_SUCCEEDED':
        try:
            if batch_method == "inline":
                # Inline results
                for i, inline_response in enumerate(batch_job.dest.inlined_responses):
                    resume_path = list(uploaded_resumes.keys())[i]
                    resume_uploaded, resume_db_file_id = uploaded_resumes[resume_path]
                    
                    if inline_response.response:
                        response_text = inline_response.response.text
                        
                        if use_structured_output:
                            try:
                                result_data = json.loads(response_text)
                                # Add automatic timestamp
                                current_time = datetime.now().isoformat() + "Z"
                                result_data["evaluation_timestamp"] = current_time
                                result_json = json.dumps(result_data, indent=2)
                                results[resume_path] = result_json
                                
                                # Store analysis result in database if username is provided
                                if username and user_id and jd_db_file_id and resume_db_file_id:
                                    # Create a cache record for this analysis
                                    cache_id = db_manager.save_cache_record(
                                        user_id=user_id,
                                        cache_name=f"batch_analysis_{i}",
                                        display_name=f"Batch Analysis of {Path(resume_path).name}",
                                        jd_file_id=jd_db_file_id,
                                        resume_file_id=resume_db_file_id,
                                        ttl=1800
                                    )
                                    
                                    # Save the analysis result
                                    db_manager.save_analysis_result(
                                        user_id=user_id,
                                        cache_id=cache_id,
                                        jd_file_id=jd_db_file_id,
                                        resume_file_id=resume_db_file_id,
                                        result_json=result_json
                                    )
                            except json.JSONDecodeError:
                                results[resume_path] = response_text
                        else:
                            results[resume_path] = response_text
                            
                            # Store non-structured analysis in database
                            if username and user_id and jd_db_file_id and resume_db_file_id:
                                # Create a cache record for this analysis
                                cache_id = db_manager.save_cache_record(
                                    user_id=user_id,
                                    cache_name=f"batch_analysis_{i}",
                                    display_name=f"Batch Analysis of {Path(resume_path).name}",
                                    jd_file_id=jd_db_file_id,
                                    resume_file_id=resume_db_file_id,
                                    ttl=1800
                                )
                                
                                # Save the analysis result
                                db_manager.save_analysis_result(
                                    user_id=user_id,
                                    cache_id=cache_id,
                                    jd_file_id=jd_db_file_id,
                                    resume_file_id=resume_db_file_id,
                                    result_json=json.dumps({"text_result": response_text})
                                )
                            
                    elif inline_response.error:
                        results[resume_path] = f"Error: {inline_response.error}"
            else:
                # File results
                result_file_name = batch_job.dest.file_name
                file_content = client.files.download(file=result_file_name)
                
                lines = file_content.decode('utf-8').strip().split('\n')
                for line in lines:
                    line_data = json.loads(line)
                    if 'key' in line_data:
                        # Extract resume index from key
                        resume_index = int(line_data['key'].split('-')[1])
                        resume_path = list(uploaded_resumes.keys())[resume_index]
                        resume_uploaded, resume_db_file_id = uploaded_resumes[resume_path]
                        
                        if 'response' in line_data:
                            response_text = line_data['response']['candidates'][0]['content']['parts'][0]['text']
                            
                            if use_structured_output:
                                try:
                                    result_data = json.loads(response_text)
                                    # Add automatic timestamp
                                    current_time = datetime.now().isoformat() + "Z"
                                    result_data["evaluation_timestamp"] = current_time
                                    result_json = json.dumps(result_data, indent=2)
                                    results[resume_path] = result_json
                                    
                                    # Store analysis result in database if username is provided
                                    if username and user_id and jd_db_file_id and resume_db_file_id:
                                        # Create a cache record for this analysis
                                        cache_id = db_manager.save_cache_record(
                                            user_id=user_id,
                                            cache_name=f"batch_analysis_{resume_index}",
                                            display_name=f"Batch Analysis of {Path(resume_path).name}",
                                            jd_file_id=jd_db_file_id,
                                            resume_file_id=resume_db_file_id,
                                            ttl=1800
                                        )
                                        
                                        # Save the analysis result
                                        db_manager.save_analysis_result(
                                            user_id=user_id,
                                            cache_id=cache_id,
                                            jd_file_id=jd_db_file_id,
                                            resume_file_id=resume_db_file_id,
                                            result_json=result_json
                                        )
                                except json.JSONDecodeError:
                                    results[resume_path] = response_text
                            else:
                                results[resume_path] = response_text
                                
                                # Store non-structured analysis in database
                                if username and user_id and jd_db_file_id and resume_db_file_id:
                                    # Create a cache record for this analysis
                                    cache_id = db_manager.save_cache_record(
                                        user_id=user_id,
                                        cache_name=f"batch_analysis_{resume_index}",
                                        display_name=f"Batch Analysis of {Path(resume_path).name}",
                                        jd_file_id=jd_db_file_id,
                                        resume_file_id=resume_db_file_id,
                                        ttl=1800
                                    )
                                    
                                    # Save the analysis result
                                    db_manager.save_analysis_result(
                                        user_id=user_id,
                                        cache_id=cache_id,
                                        jd_file_id=jd_db_file_id,
                                        resume_file_id=resume_db_file_id,
                                        result_json=json.dumps({"text_result": response_text})
                                    )
                        elif 'error' in line_data:
                            results[resume_path] = f"Error: {line_data['error']}"
        except Exception as e:
            print(f"Error retrieving results: {e}")
            # For any resume paths not already in results
            for resume_path in uploaded_resumes.keys():
                if resume_path not in results:
                    results[resume_path] = f"Error retrieving result: {str(e)}"
    else:
        print(f"Batch job failed with state: {batch_job.state.name}")
        if hasattr(batch_job, 'error'):
            print(f"Error: {batch_job.error}")
            
        # Add error information to results for all resumes
        for resume_path in uploaded_resumes.keys():
            error_message = f"Batch job failed with state: {batch_job.state.name}"
            if hasattr(batch_job, 'error'):
                error_message += f", Error: {batch_job.error}"
            results[resume_path] = error_message
    
    # Clean up uploaded files
    print("Cleaning up uploaded files...")
    try:
        client.files.delete(name=jd_uploaded.name)
        for uploaded, _ in uploaded_resumes.values():
            client.files.delete(name=uploaded.name)
    except Exception as e:
        print(f"Warning: Could not clean up some files: {e}")
    
    return results


def analyze_bulk_resumes_parallel(jd_file_path, resume_file_paths, username=None, use_structured_output=True, max_workers=5):
    """
    Analyze multiple resumes against a single job description in parallel without using Batch API.
    
    Args:
        jd_file_path (str): Path to the job description file (.pdf or .docx)
        resume_file_paths (list): List of paths to resume files (.pdf or .docx)
        username (str): Username to associate with the files and analysis results
        use_structured_output (bool): Whether to use structured output with ATS schema
        max_workers (int): Maximum number of parallel workers (threads)
        
    Returns:
        dict: Dictionary with resume paths as keys and analysis results as values
    """
    # Helper functions
    def _validate_file_type(path: Path):
        if path.suffix.lower() not in {".pdf", ".docx"}:
            raise ValueError(f"Unsupported file type: {path.suffix}. Only PDF and DOCX are supported.")

    def _get_mime_type(path: Path) -> str:
        return "application/pdf" if path.suffix.lower() == ".pdf" else "text/plain"

    def _convert_docx_to_text(path: Path) -> str:
        from docx import Document
        doc = Document(path)
        return "\n".join(p.text for p in doc.paragraphs)
    
    # Function to process a single resume against the job description
    def process_resume(resume_path, jd_file_path, username, use_structured_output):
        try:
            print(f"Processing resume: {Path(resume_path).name}")
            # Create a new client instance per thread to avoid thread safety issues
            client = genai.Client(api_key=GEMINI_API_KEY)
            
            # Get or create user if username is provided
            user_id = None
            if username:
                user = db_manager.get_or_create_user(username)
                user_id = user['id']
            
            # Prepare resume file
            resume_path_obj = Path(resume_path)
            _validate_file_type(resume_path_obj)
            resume_file_type = resume_path_obj.suffix.lower()[1:]  # Remove the dot
            
            if resume_path_obj.suffix.lower() == ".docx":
                resume_text = _convert_docx_to_text(resume_path_obj)
                resume_io = io.BytesIO(resume_text.encode("utf-8"))
                resume_mime_type = "text/plain"
            else:
                resume_io = io.BytesIO(resume_path_obj.read_bytes())
                resume_mime_type = _get_mime_type(resume_path_obj)
            
            # Prepare JD file
            jd_path_obj = Path(jd_file_path)
            _validate_file_type(jd_path_obj)
            jd_file_type = jd_path_obj.suffix.lower()[1:]  # Remove the dot
            
            if jd_path_obj.suffix.lower() == ".docx":
                jd_text = _convert_docx_to_text(jd_path_obj)
                jd_io = io.BytesIO(jd_text.encode("utf-8"))
                jd_mime_type = "text/plain"
            else:
                jd_io = io.BytesIO(jd_path_obj.read_bytes())
                jd_mime_type = _get_mime_type(jd_path_obj)
            
            # Upload files
            resume_uploaded, resume_db_file_id = upload_file_with_retry(
                client, 
                resume_io, 
                resume_mime_type,
                username=username,
                filename=resume_path_obj.name,
                file_type=resume_file_type
            )
            
            jd_uploaded, jd_db_file_id = upload_file_with_retry(
                client, 
                jd_io, 
                jd_mime_type,
                username=username,
                filename=jd_path_obj.name,
                file_type=jd_file_type
            )
            
            # Wait if processing
            while hasattr(resume_uploaded, "state") and resume_uploaded.state.name == "PROCESSING":
                time.sleep(2)
                resume_uploaded = client.files.get(name=resume_uploaded.name)
                
            while hasattr(jd_uploaded, "state") and jd_uploaded.state.name == "PROCESSING":
                time.sleep(2)
                jd_uploaded = client.files.get(name=jd_uploaded.name)
            
            # Create the analysis request
            analysis_prompt = f"""
{prompts.system_prompt}

Please analyze the following resume against the job description and provide a comprehensive ATS evaluation.

Job Description: (see attached file)
Resume: (see attached file)
"""
            
            # Create cache with both files
            display_name = f"Cache with {jd_path_obj.name} and {resume_path_obj.name}"
            cache = client.caches.create(
                model=GEMINI_MODEL,
                config=types.CreateCachedContentConfig(
                    display_name=display_name,
                    system_instruction=prompts.system_prompt,
                    contents=[jd_uploaded, resume_uploaded],
                    ttl="1800s",
                ),
            )
            
            # Store cache information in the database if username is provided
            cache_id = None
            
            if username and user_id and jd_db_file_id and resume_db_file_id:
                cache_id = db_manager.save_cache_record(
                    user_id=user_id,
                    cache_name=cache.name,
                    display_name=display_name,
                    jd_file_id=jd_db_file_id,
                    resume_file_id=resume_db_file_id,
                    ttl=1800
                )
            
            # Generate content
            if use_structured_output:
                response = client.models.generate_content(
                    model=GEMINI_MODEL,
                    contents="Analyze the resume against the job description.",
                    config=types.GenerateContentConfig(
                        cached_content=cache.name,
                        response_schema=ATS_SCHEMA,
                        response_mime_type="application/json"
                    ),
                )
                
                # Parse JSON response
                try:
                    result_data = json.loads(response.text)
                    current_time = datetime.now().isoformat() + "Z"
                    result_data["evaluation_timestamp"] = current_time
                    result_json = json.dumps(result_data, indent=2)
                    
                    # Store the analysis result in the database if username is provided
                    if username and user_id and cache_id and jd_db_file_id and resume_db_file_id:
                        db_manager.save_analysis_result(
                            user_id=user_id,
                            cache_id=cache_id,
                            jd_file_id=jd_db_file_id,
                            resume_file_id=resume_db_file_id,
                            result_json=result_json
                        )
                    
                    result = result_json
                except json.JSONDecodeError:
                    result = response.text
            else:
                response = client.models.generate_content(
                    model=GEMINI_MODEL,
                    contents=analysis_prompt,
                    config=types.GenerateContentConfig(cached_content=cache.name),
                )
                result_text = response.text
                
                # Store the analysis result in the database if username is provided
                if username and user_id and cache_id and jd_db_file_id and resume_db_file_id:
                    db_manager.save_analysis_result(
                        user_id=user_id,
                        cache_id=cache_id,
                        jd_file_id=jd_db_file_id,
                        resume_file_id=resume_db_file_id,
                        result_json=json.dumps({"text_result": result_text})
                    )
                
                result = result_text
            
            # Clean up uploaded files
            try:
                client.files.delete(name=resume_uploaded.name)
                client.files.delete(name=jd_uploaded.name)
            except Exception as e:
                print(f"Warning: Could not clean up files for {resume_path_obj.name}: {e}")
            
            print(f"Completed: {resume_path_obj.name}")
            return resume_path, result
            
        except Exception as e:
            print(f"Error processing {resume_path}: {e}")
            return resume_path, f"Error: {str(e)}"
    
    # Print start information
    print(f"Starting parallel analysis of {len(resume_file_paths)} resumes...")
    print(f"Maximum parallel workers: {max_workers}")
    
    # Create a progress counter with thread safety
    progress_lock = threading.Lock()
    completed_count = 0
    total_count = len(resume_file_paths)
    
    def update_progress(future):
        nonlocal completed_count
        with progress_lock:
            completed_count += 1
            print(f"Progress: {completed_count}/{total_count} resumes processed ({completed_count/total_count*100:.1f}%)")
    
    # Process resumes in parallel
    start_time = time.time()
    results = {}
    
    with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
        # Submit all tasks
        future_to_resume = {
            executor.submit(
                process_resume, resume_path, jd_file_path, username, use_structured_output
            ): resume_path for resume_path in resume_file_paths
        }
        
        # Add progress tracking callback
        for future in future_to_resume:
            future.add_done_callback(update_progress)
        
        # Collect results as they complete
        for future in concurrent.futures.as_completed(future_to_resume):
            try:
                resume_path, result = future.result()
                results[resume_path] = result
            except Exception as e:
                resume_path = future_to_resume[future]
                results[resume_path] = f"Unexpected error: {str(e)}"
    
    end_time = time.time()
    print(f"Parallel processing completed in {end_time - start_time:.2f} seconds")
    
    return results


# Example code for parallel resume analysis with all sample resumes
# To run this code, uncomment it and run this module directly
if __name__ == "__main__":
    parallel_results = analyze_bulk_resumes_parallel(
        "D:\\Coding\\Context-cache-system\\samples\\sample_jd_1.pdf",
        [
            "D:\\Coding\\Context-cache-system\\samples\\resume_1.pdf",
            "D:\\Coding\\Context-cache-system\\samples\\Resume_10.pdf",
            "D:\\Coding\\Context-cache-system\\samples\\RESUME_2.pdf",
            "D:\\Coding\\Context-cache-system\\samples\\resume_3.pdf",
            "D:\\Coding\\Context-cache-system\\samples\\Resume_4.pdf",
            "D:\\Coding\\Context-cache-system\\samples\\Resume_5.pdf",
            "D:\\Coding\\Context-cache-system\\samples\\Resume_6.pdf",
            "D:\\Coding\\Context-cache-system\\samples\\Resume_7.pdf",
            "D:\\Coding\\Context-cache-system\\samples\\Resume_8.pdf",
            "D:\\Coding\\Context-cache-system\\samples\\Resume_9.pdf",
        ],
        use_structured_output=True,
        max_workers=5
    )

    print("\n=== Parallel Resume Analysis Results ===")

    # Create a dictionary to store all results
    results = {"results": []}

    for resume_path, result in parallel_results.items():
        # Print condensed result without file name
        print("\n--- Analysis Result ---")
        print(result)
        
        # Parse the JSON string and add to the results list
        try:
            result_data = json.loads(result)
            results["results"].append(result_data)
        except json.JSONDecodeError:
            # If not valid JSON, store as string
            results["results"].append({"error": result})

    # Save all results to a single JSON file
    with open('data.json', 'w') as f:
        json.dump(results, f, indent=2)

    print(f"\nAll results saved to data.json")

# Uncomment to run the batch API analysis with all sample resumes

# bulk_results = analyze_bulk_resumes(
#     "D:\\Coding\\Context-cache-system\\samples\\sample_jd_1.pdf",
#     [
#         "D:\\Coding\\Context-cache-system\\samples\\resume_1.pdf",
#         "D:\\Coding\\Context-cache-system\\samples\\Resume_10.pdf",
#         "D:\\Coding\\Context-cache-system\\samples\\RESUME_2.pdf",
#         "D:\\Coding\\Context-cache-system\\samples\\resume_3.pdf",
#         "D:\\Coding\\Context-cache-system\\samples\\Resume_4.pdf",
#         "D:\\Coding\\Context-cache-system\\samples\\Resume_5.pdf",
#         "D:\\Coding\\Context-cache-system\\samples\\Resume_6.pdf",
#         "D:\\Coding\\Context-cache-system\\samples\\Resume_7.pdf",
#         "D:\\Coding\\Context-cache-system\\samples\\Resume_8.pdf",
#         "D:\\Coding\\Context-cache-system\\samples\\Resume_9.pdf",
#     ],
#     use_structured_output=True,
#     batch_method="inline"  # Use "file" for large batches (>50 resumes)
# )

# print("\\n=== Bulk Resume Analysis Results ===")
# # Create a dictionary to store all results
# combined_results = {"results": []}

# for resume_path, result in bulk_results.items():
#     # Print condensed result without file name
#     print("\\n--- Analysis Result ---")
#     print(result)
    
#     # Parse the JSON string and add to the results list
#     try:
#         result_data = json.loads(result)
#         combined_results["results"].append(result_data)
#     except json.JSONDecodeError:
#         # If not valid JSON, store as string
#         combined_results["results"].append({"error": result})

# # Save all results to a single JSON file
# with open('data.json', 'w') as f:
#     json.dump(combined_results, indent=2, fp=f)
        
# print(f"\\nAll results saved to data.json")

# Example usage:
# For single resume analysis
# """
# single_result = analyze_two_files(
#     "path/to/resume.pdf",
#     "path/to/job_description.pdf",
#     use_structured_output=True
# )
# print(single_result)
# """

# # Example of batch analysis with Batch API
# """
# # Test with a smaller batch first (2-3 resumes)
# bulk_results = analyze_bulk_resumes(
#     "path/to/job_description.pdf",
#     [
#         "path/to/resume1.pdf",
#         "path/to/resume2.pdf",
#         "path/to/resume3.pdf",
#     ],
#     use_structured_output=True,
#     batch_method="inline"  # Use "inline" for small batches, "file" for large batches (>50 resumes)
# )

# # Process and display results
# for resume_path, result in bulk_results.items():
#     print(f"\\n--- {Path(resume_path).name} ---")
#     # Optional: Save result to a file
#     # with open(f"{Path(resume_path).stem}_analysis.json", "w") as f:
#     #     f.write(result)
#     print(result)
# """

# # Example of parallel processing without Batch API
# """
# # Process multiple resumes in parallel
# parallel_results = analyze_bulk_resumes_parallel(
#     "path/to/job_description.pdf",
#     [
#         "path/to/resume1.pdf",
#         "path/to/resume2.pdf",
#         "path/to/resume3.pdf",
#     ],
#     use_structured_output=True,
#     max_workers=3  # Adjust based on your system's capabilities
# )

# # Process and display results
# for resume_path, result in parallel_results.items():
#     print(f"\\n--- {Path(resume_path).name} ---")
#     print(result)
# """

